from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "recovery" / "codex_exact_copy.docx"
SECTION_MD = ROOT / "docs" / "project_structure_section_he.md"
OUTPUT_DOC = ROOT / "recovery" / "book_with_merged_project_structure.docx"

START_HEADING = "מבנה הפרויקט - קוד הפרויקט"
END_HEADING = "למידת חיזוקים Reinforcement Learning"


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_before(ref: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref._element.addprevious(new_p)
    paragraph = Paragraph(new_p, ref._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def set_run_font(paragraph: Paragraph, name: str = "Arial", size_pt: int = 12) -> None:
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)


def add_code_block_before(ref: Paragraph, lines: list[str]) -> None:
    for line in lines:
        p = insert_paragraph_before(ref, line, style="No Spacing")
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    return text


def parse_markdown_lines(lines: list[str]) -> list[tuple[str, list[str] | str]]:
    items: list[tuple[str, list[str] | str]] = []
    in_code = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.startswith("```"):
            if in_code:
                items.append(("code", code_lines[:]))
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            items.append(("blank", ""))
            continue

        if line.startswith("# "):
            items.append(("h1", clean_inline(line[2:].strip())))
            continue
        if line.startswith("## "):
            items.append(("h2", clean_inline(line[3:].strip())))
            continue
        if line.startswith("### "):
            items.append(("h3", clean_inline(line[4:].strip())))
            continue
        if re.match(r"^\d+\.\s+", line):
            items.append(("num", clean_inline(line)))
            continue
        if line.startswith("- "):
            items.append(("bullet", clean_inline(line[2:].strip())))
            continue

        items.append(("p", clean_inline(line)))

    return items


def replace_section(doc: Document, items: list[tuple[str, list[str] | str]]) -> None:
    start_idx = None
    end_idx = None

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == START_HEADING and start_idx is None:
            start_idx = idx
        if text == END_HEADING and start_idx is not None:
            end_idx = idx
            break

    if start_idx is None or end_idx is None or end_idx <= start_idx:
        raise RuntimeError("Could not locate project structure chapter boundaries.")

    end_paragraph = doc.paragraphs[end_idx]

    for paragraph in list(doc.paragraphs[start_idx:end_idx]):
        delete_paragraph(paragraph)

    previous_blank = True
    for kind, payload in items:
        if kind == "blank":
            if not previous_blank:
                insert_paragraph_before(end_paragraph, "")
            previous_blank = True
            continue

        previous_blank = False
        if kind == "code":
            add_code_block_before(end_paragraph, payload)  # type: ignore[arg-type]
            continue

        style = "Normal"
        text = payload  # type: ignore[assignment]
        if kind == "h1":
            style = "Heading 1"
        elif kind == "h2":
            style = "Heading 2"
        elif kind == "h3":
            style = "Heading 3"
        elif kind == "bullet":
            style = "Normal"
            text = f"- {text}"
        elif kind == "num":
            style = "Normal"

        paragraph = insert_paragraph_before(end_paragraph, str(text), style=style)
        if style == "Normal":
            set_run_font(paragraph, "Arial", 12)


def main() -> None:
    doc = Document(SOURCE_DOC)
    lines = SECTION_MD.read_text(encoding="utf-8").splitlines()
    items = parse_markdown_lines(lines)
    replace_section(doc, items)
    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()

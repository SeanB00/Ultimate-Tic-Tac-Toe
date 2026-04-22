from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
START_HEADING = "\u05e7\u05d5\u05d3 \u05de\u05dc\u05d0 \u05e9\u05dc \u05d4\u05e4\u05e8\u05d5\u05d9\u05e7\u05d8"
END_HEADING = "\u05d1\u05d9\u05d1\u05dc\u05d9\u05d5\u05d2\u05e8\u05e4\u05d9\u05d4"
INTRO_TEXT = (
    "\u05dc\u05d4\u05dc\u05df \u05e7\u05d1\u05e6\u05d9 \u05e7\u05d5\u05d3 \u05d4\u05de\u05e7\u05d5\u05e8 "
    "\u05d4\u05de\u05e2\u05d5\u05d3\u05db\u05e0\u05d9\u05dd \u05e9\u05dc \u05d4\u05e4\u05e8\u05d5\u05d9\u05e7\u05d8 "
    "\u05db\u05e4\u05d9 \u05e9\u05d4\u05dd \u05e7\u05d9\u05d9\u05de\u05d9\u05dd \u05db\u05e2\u05ea \u05d1\u05de\u05d0\u05d2\u05e8 "
    "\u05d4\u05e7\u05d5\u05d3. \u05e7\u05d1\u05e6\u05d9 \u05d3\u05d0\u05d8\u05d4, \u05de\u05d5\u05d3\u05dc\u05d9\u05dd "
    "\u05d1\u05d9\u05e0\u05d0\u05e8\u05d9\u05d9\u05dd, \u05d5\u05ea\u05d9\u05e7\u05d9\u05d5\u05ea \u05e1\u05d1\u05d9\u05d1\u05ea "
    "\u05e4\u05d9\u05ea\u05d5\u05d7 \u05dc\u05d0 \u05e0\u05db\u05dc\u05dc\u05d5, \u05db\u05d9 \u05d4\u05dd "
    "\u05d0\u05d9\u05e0\u05dd \u05e7\u05d5\u05d3 \u05de\u05e7\u05d5\u05e8."
)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_before(
    ref: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref._element.addprevious(new_p)
    paragraph = Paragraph(new_p, ref._parent)
    if style is not None:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def set_run_font(run, name: str = "Arial", size_pt: float = 12.0, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)


def set_paragraph_rtl(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)


def set_paragraph_ltr(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is not None:
        p_pr.remove(bidi)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def tracked_source_files() -> list[str]:
    output = subprocess.check_output(
        ["git", "ls-files", "requirements.txt", "*.py"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
    )
    return [line for line in output.splitlines() if line.strip()]


def sanitize_code(text: str) -> str:
    text = re.sub(
        r'API_KEY\s*=\s*["\'][^"\']+["\']',
        'API_KEY = "<API_KEY>"',
        text,
    )
    return text.lstrip("\ufeff").replace("\r\n", "\n")


def add_heading_before(ref: Paragraph, text: str, level: int) -> Paragraph:
    heading = insert_paragraph_before(ref, text, style=f"Heading {level}")
    if any("\u0590" <= ch <= "\u05ff" for ch in text):
        set_paragraph_rtl(heading)
    else:
        set_paragraph_ltr(heading)
    for run in heading.runs:
        set_run_font(run, size_pt=14 if level == 2 else 16, bold=True)
    return heading


def add_body_before(ref: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_before(ref, text, style="Normal")
    set_paragraph_rtl(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size_pt=12)
    return paragraph


def add_code_before(ref: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_before(ref, style="Normal")
    set_paragraph_ltr(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, name="Courier New", size_pt=7.5)
    return paragraph


def replace_full_code_section(doc: Document, files: list[str]) -> None:
    start_idx = None
    end_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == START_HEADING and start_idx is None:
            start_idx = idx
        elif text == END_HEADING and start_idx is not None:
            end_idx = idx
            break

    if start_idx is None or end_idx is None or end_idx <= start_idx:
        raise RuntimeError("Could not locate the full-code section boundaries in the document.")

    end_paragraph = doc.paragraphs[end_idx]

    for paragraph in list(doc.paragraphs[start_idx + 1:end_idx]):
        delete_paragraph(paragraph)

    add_body_before(end_paragraph, INTRO_TEXT)
    insert_paragraph_before(end_paragraph, "")

    for rel_path in files:
        add_heading_before(end_paragraph, rel_path, level=2)
        code_text = sanitize_code((ROOT / rel_path).read_text(encoding="utf-8", errors="replace"))
        add_code_before(end_paragraph, code_text)


def save_document(doc: Document, doc_path: Path) -> Path:
    try:
        doc.save(doc_path)
        return doc_path
    except PermissionError:
        fallback = doc_path.with_name(f"{doc_path.stem} - full code updated{doc_path.suffix}")
        doc.save(fallback)
        return fallback


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Refresh the 'full code' chapter in the UTTT project book.",
    )
    parser.add_argument("docx_path", help="Path to the .docx file to update.")
    return parser.parse_args()


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    args = parse_args()
    doc_path = Path(args.docx_path).expanduser().resolve()
    doc = Document(doc_path)
    set_update_fields_on_open(doc)
    replace_full_code_section(doc, tracked_source_files())
    output_path = save_document(doc, doc_path)
    print(output_path)


if __name__ == "__main__":
    main()
